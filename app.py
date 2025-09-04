# app.py - Веб-сервер для обработки DOC/DOCX файлов
from flask import Flask, request, jsonify
from flask_cors import CORS
import io
import base64
import logging
from datetime import datetime
import os
import tempfile
import traceback

# Импорт библиотек для работы с документами
try:
    from docx import Document
    import mammoth
    import python_docx_replace as docx_replace
except ImportError:
    pass

app = Flask(__name__)
CORS(app)  # Разрешаем CORS для работы с n8n

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Класс для обработки DOC/DOCX документов"""
    
    @staticmethod
    def extract_text_from_docx(file_content):
        """Извлечение текста из DOCX с сохранением структуры"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Собираем весь текст с сохранением структуры
            full_text = []
            
            # Обработка параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Обработка таблиц
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    full_text.append("\n[ТАБЛИЦА]\n" + "\n".join(table_text) + "\n[/ТАБЛИЦА]")
            
            return "\n\n".join(full_text)
        except Exception as e:
            logger.error(f"Ошибка при обработке DOCX: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_doc(file_content):
        """Извлечение текста из DOC файла используя mammoth"""
        try:
            # Сохраняем во временный файл для mammoth
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            # Конвертируем DOC в текст
            with open(tmp_file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value
            
            # Удаляем временный файл
            os.unlink(tmp_file_path)
            
            return text
        except Exception as e:
            logger.error(f"Ошибка при обработке DOC: {str(e)}")
            # Пробуем альтернативный метод
            return DocumentProcessor.extract_text_simple(file_content)
    
    @staticmethod
    def extract_text_simple(file_content):
        """Простое извлечение текста как fallback"""
        try:
            # Попытка декодировать как текст
            text = file_content.decode('utf-8', errors='ignore')
            # Очистка от бинарных символов
            text = ''.join(char for char in text if char.isprintable() or char.isspace())
            return text
        except:
            return "Не удалось извлечь текст из документа"
    
    @staticmethod
    def analyze_contract_structure(text):
        """Анализ структуры договора для ИИ-юриста"""
        structure = {
            "has_parties": False,
            "has_subject": False,
            "has_terms": False,
            "has_responsibilities": False,
            "has_signatures": False,
            "sections": []
        }
        
        # Ключевые слова для определения разделов
        keywords = {
            "parties": ["стороны", "заказчик", "исполнитель", "продавец", "покупатель", "арендатор", "арендодатель"],
            "subject": ["предмет договора", "предмет соглашения"],
            "terms": ["сроки", "срок действия", "период"],
            "responsibilities": ["обязанности", "ответственность", "обязательства"],
            "signatures": ["подписи сторон", "реквизиты"]
        }
        
        text_lower = text.lower()
        
        # Проверка наличия ключевых разделов
        for key, words in keywords.items():
            for word in words:
                if word in text_lower:
                    structure[f"has_{key}"] = True
                    break
        
        # Извлечение заголовков разделов
        lines = text.split('\n')
        for line in lines:
            line_stripped = line.strip()
            if line_stripped and (
                line_stripped.isupper() or 
                any(char.isdigit() for char in line_stripped[:3]) or
                line_stripped.startswith('§') or
                line_stripped.startswith('Статья')
            ):
                if len(line_stripped) < 100:  # Ограничиваем длину заголовка
                    structure["sections"].append(line_stripped)
        
        return structure

@app.route('/health', methods=['GET'])
def health_check():
    """Проверка работоспособности сервера"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "service": "Document Processor for n8n"
    })

@app.route('/process', methods=['POST'])
def process_document():
    """
    Основной эндпоинт для обработки документов
    Принимает файл в base64 или бинарном формате
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "error": "Нет данных в запросе"
            }), 400
        
        # Получаем содержимое файла
        file_content = None
        file_name = data.get('filename', 'document.docx')
        
        # Проверяем разные форматы входных данных
        if 'file_base64' in data:
            # Файл в base64
            file_content = base64.b64decode(data['file_base64'])
        elif 'file_content' in data:
            # Прямое содержимое файла
            if isinstance(data['file_content'], str):
                file_content = data['file_content'].encode('utf-8')
            else:
                file_content = data['file_content']
        elif 'binary' in data:
            # Бинарные данные из n8n
            file_content = base64.b64decode(data['binary'])
        else:
            return jsonify({
                "success": False,
                "error": "Файл не найден в запросе. Используйте поля: file_base64, file_content или binary"
            }), 400
        
        # Определяем тип файла и извлекаем текст
        processor = DocumentProcessor()
        
        if file_name.endswith('.docx'):
            text = processor.extract_text_from_docx(file_content)
        elif file_name.endswith('.doc'):
            text = processor.extract_text_from_doc(file_content)
        else:
            # Пробуем как DOCX по умолчанию
            try:
                text = processor.extract_text_from_docx(file_content)
            except:
                text = processor.extract_text_from_doc(file_content)
        
        # Анализируем структуру договора
        structure = processor.analyze_contract_structure(text)
        
        # Подготавливаем ответ
        response = {
            "success": True,
            "filename": file_name,
            "text": text,
            "text_length": len(text),
            "word_count": len(text.split()),
            "contract_analysis": structure,
            "metadata": {
                "processed_at": datetime.now().isoformat(),
                "file_size_bytes": len(file_content)
            }
        }
        
        # Добавляем рекомендации для ИИ-юриста
        if text:
            response["ai_instructions"] = {
                "prompt_suggestion": f"""Проанализируй следующий договор и предоставь юридическую оценку:

1. Проверь полноту договора (наличие всех обязательных разделов)
2. Выяви потенциальные риски для клиента
3. Укажи на неоднозначные формулировки
4. Предложи улучшения
5. Оцени соответствие законодательству РФ

Текст договора:
{text[:2000]}...""",
                "has_content": True
            }
        else:
            response["ai_instructions"] = {
                "prompt_suggestion": "Документ не содержит текста для анализа",
                "has_content": False
            }
        
        logger.info(f"Успешно обработан документ: {file_name}")
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Ошибка при обработке документа: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }), 500

@app.route('/process-n8n', methods=['POST'])
def process_n8n_format():
    """
    Специальный эндпоинт для n8n с упрощенным форматом
    """
    try:
        # n8n может отправлять данные в разных форматах
        if request.content_type == 'application/json':
            data = request.get_json()
        else:
            # Если n8n отправляет файл напрямую
            file_content = request.data
            
        # Обработка для n8n Telegram node
        if isinstance(data, dict) and 'data' in data:
            # Формат от Telegram Get File
            telegram_data = data['data']
            if 'file_path' in telegram_data:
                # Здесь должен быть base64 контент файла
                file_content = base64.b64decode(telegram_data.get('file_content', ''))
            else:
                file_content = base64.b64decode(data.get('binary', ''))
        
        # Извлекаем текст
        processor = DocumentProcessor()
        text = processor.extract_text_from_docx(file_content)
        
        # Упрощенный ответ для n8n
        return jsonify({
            "text": text,
            "success": True,
            "word_count": len(text.split()),
            "ready_for_ai": True
        })
        
    except Exception as e:
        logger.error(f"Ошибка n8n обработки: {str(e)}")
        return jsonify({
            "text": "",
            "success": False,
            "error": str(e)
        }), 500

@app.route('/', methods=['GET'])
def index():
    """Главная страница с документацией"""
    return """
    <html>
    <head>
        <title>Document Processor API для n8n</title>
        <style>
            body { font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; }
            h1 { color: #333; }
            code { background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }
            pre { background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }
            .endpoint { background: #e7f3ff; padding: 10px; margin: 10px 0; border-radius: 5px; }
        </style>
    </head>
    <body>
        <h1>📄 Document Processor API для n8n</h1>
        <p>Сервис для извлечения текста из DOC/DOCX файлов для ИИ-юриста</p>
        
        <h2>Эндпоинты:</h2>
        
        <div class="endpoint">
            <h3>POST /process</h3>
            <p>Основной эндпоинт для обработки документов</p>
            <pre>
{
    "file_base64": "base64_encoded_file_content",
    "filename": "contract.docx"
}
            </pre>
        </div>
        
        <div class="endpoint">
            <h3>POST /process-n8n</h3>
            <p>Специальный эндпоинт для n8n (упрощенный формат)</p>
        </div>
        
        <div class="endpoint">
            <h3>GET /health</h3>
            <p>Проверка работоспособности сервиса</p>
        </div>
        
        <h2>Использование в n8n:</h2>
        <ol>
            <li>Telegram Trigger (получение файла)</li>
            <li>HTTP Request node → POST на /process-n8n</li>
            <li>AI Agent node (анализ текста)</li>
            <li>Telegram Send Message (ответ)</li>
        </ol>
        
        <p><strong>Статус:</strong> <span style="color: green;">✅ Сервис работает</span></p>
    </body>
    </html>
    """

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)